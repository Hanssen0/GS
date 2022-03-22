import jieba
import fitz
import hashlib
import jieba.posseg as pseg
from fitz import Rect

from init_utils import *
root = init_by_key('root')
pro_dic = {}
from zipfile import ZipFile
from bs4 import BeautifulSoup

import os,shutil

from docx import Document
# root = '/home/oo/wordCl/'
root = init_by_key('root')

def readpdf(pdf_file):
    pdf = None
    text = ''
    try:
        pdf=fitz.open(pdf_file)
    except:
        return text
    heads= []
    # table = [{}*5][]

    article_inf = {}
    blocks_list = []
    for index,page in enumerate(pdf):
        # img = get_img(file, index)   #pdf转图
        # blocks = page.getTextWords()
        text1=page.getText().replace( '\n' , '' )

        text+=text1
    return text
def readdocx_zip(docx_file):
    text = ''
    try:
        document=ZipFile(docx_file)
        xml=document.read("word/document.xml")
        wordObj=BeautifulSoup(xml.decode("utf-8"))
        doc=wordObj.findAll("w:t")

    except:
        return text
    for t in doc:
        t_text = t.text.strip().replace('\n','')
        if(t_text not in text):
            text+=t_text
    return text


def readdocx(docx_file):

    text = readdocx_org(docx_file)
    if(text == ''):
        text = readdocx_zip(docx_file)
    return text
def readdocx_org(docx_file):
    doc = None
    text = ''
    try:
        doc = Document(docx_file)
    except:
        return text

    #每一段的内容
    for para in doc.paragraphs:
        p_t = para.text.replace('\n','')
        if(p_t not in text):

            text=text+p_t
            if para.text.endswith('。'):
                text += '。'


    # #每一段的编号、内容
    # for i in range(len(doc.paragraphs)):
    #     print(str(i),  doc.paragraphs[i].text)

    #表格
    tbs = doc.tables
    for tb in tbs:
        #行
        for row in tb.rows:
            #列
            for cell in row.cells:
                c_text = cell.text.replace( '\n' , '' )
                if ( c_text not in text):
                    text += c_text
                    # print(cell.text)
                #也可以用下面方法
                '''text = ''
                for p in cell.paragraphs:
                    text += p.text
                print(text)'''
    return text

#shutil.copyfile应用,参数必须具体到文件名
def copyfile(srcfile,dstfile):
    if not os.path.isfile(srcfile):
        print("%s not exit!" % (srcfile))
    else:
        fpath,fname=os.path.split(dstfile)
    if not os.path.exists(fpath):
        os.makedirs(fpath)
    shutil.copyfile(srcfile,dstfile)
    #print("copy %s" % (srcfile,dstfile))


def listdir(path ,list = []):

    for file in os.listdir(path):
        file_path = os.path.join(path, file)
        if os.path.isdir(file_path):
            listdir(file_path, list)
        else:
            list.append(file_path)
    return list




# 创建停用词list
def stopwordslist(filepath):
    stopwords = [line.strip() for line in open(filepath, 'r', encoding='utf-8').readlines()]
    return stopwords

def use_jieba_parti(sen = ''):
    nouns = set()

    jieba.load_userdict(root+'data/refine_cla/dict/all_coll.txt')
    words = pseg.cut(sen)
    stopwords = stopwordslist(root+'data/stopwords.txt')
    for word, flag in words:
        if word not in stopwords:
            if('n' in flag):
                nouns.add(word)
    return nouns
                # print ('%s %s'%  (word, flag))

def initdict():
    dic = {}
    dic['__label__ACADEMIC'] = inititem()
    dic['__label__INDUSTRIAL'] = inititem()
    dic['__label__BUSINESS'] = inititem()
    dic['__label__INVESTMENT'] =inititem()
    dic['__label__PATENTS'] = inititem()
    return dic
def inititem():
    dic = {}
    dic['items'] = []
    dic['score'] = 0
    return dic

def init_pro_dic():
    global pro_dic
    f_path = root+'data/level_E/t_grade.txt'
    with open(f_path,'r') as  txt_file:
        lines = txt_file.readlines()
        for l in lines:
            ls = l.split(';')
            name = ls[0].strip()
            sc = int(ls[-1].strip())
            pro_dic[name]=sc
    return pro_dic








def judge_score(project_dic,comp = False):
    ig_list = ['NAME']
    dic = {}
    score = 60
    for key in project_dic.keys():
        if(not key in ig_list):
            if(type(project_dic[key])==dict):
                for item in project_dic[key].keys():
                    score+=project_dic[key][item]['score']
            else:
                for item in project_dic[key]:
                    score+=item['score']
    dic['AI'] = score
    if(comp):
        global pro_dic
        if(len(pro_dic)==0):
            pro_dic = init_pro_dic()
        if project_dic['NAME'] in pro_dic.keys():
            dic['TEACHER'] = pro_dic[project_dic['NAME']]
    return dic




def judge_score_list(name,pro,comp = False):
    ig_list = ['NAME']
    dic = {}
    score = 60
    for project_dic in pro:
        for key in project_dic.keys():
            if(not key in ig_list):
                if(type(project_dic[key])==dict):
                    for item in project_dic[key].keys():
                        score+=project_dic[key][item]['score']
                else:
                    for item in project_dic[key]:
                        score+=item['score']
    if(score)>100:
        score = 98
    dic['AI'] = score
    if(comp):
        global pro_dic
        if(len(pro_dic)==0):
            pro_dic = init_pro_dic()
        if name in pro_dic.keys():
            dic['TRUES'] = pro_dic[name]
            dic['BIAS'] = dic['AI']- dic['TRUES']
    return dic


def hash_file(filename):
    sha256_hash = hashlib.sha256()
    with open(filename,"rb") as f:
        for byte_block in iter(lambda: f.read(4096),b""):
            sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

def get_fname_by_path(f_path):
    tmp  = f_path.split('/')
    res = os.path.splitext(tmp[-1])
    return res[0]

def get_ftype_by_path(f_path):
    tmp  = f_path.split('/')
    res = os.path.splitext(tmp[-1])
    return res[-1]

def hash_list(folder_path):
    list = listdir(folder_path,[])

    h_l = []
    h_txt = ''
    for filename in list:
        h_l.append(hash_file(filename))
    h_l.sort()
    for h in h_l:
        h_txt+=str(h)
    return h_txt
def pdf_hl(f_path,l,js):
    try:
        pdf=fitz.open(f_path)
    except:
        return
    heads= []
    # table = [{}*5][]

    article_inf = {}
    blocks_list = []
    index = 0
    for index,page in enumerate(pdf):
        text=page.getText()

        blocks = page.get_text("blocks")
        lines = page.get_text("lines")
        # for blo in blocks:
        #     for item in l :
        #         if(item in blo[4]):
        #             # for txt in page.searchFor(blo[4]):
        #             page.addHighlightAnnot(Rect(blo[0],blo[1],blo[2],blo[3]))

        if('。' in text):
            sens=text.replace('\n',' ').split('。')

            for blo in sens:
                for item in l :
                    if(item in blo):
                        for txt in page.searchFor(blo.strip()):
                            highlight = page.addHighlightAnnot(txt)
                            highlight.setColors({"stroke":(0, 1, 0), "fill":(0.75, 0.8, 0.95)})
                            highlight.update()

        else:
            if(index==0):
                continue
            sens=text.split('\n')

            for blo in sens:
                for item in l :
                    if(item in blo):
                        for txt in page.searchFor(blo.strip()):
                            highlight = page.addHighlightAnnot(txt)
                            # highlight.setColors({"stroke":(1, 0, 0), "fill":(0.75, 0.8, 0.95)})
                            #
                            # highlight.update()

    if(index>3):
        txt = ""
        ch_l = {'__label__ACADEMIC':'优化文书中该项工在人才团队方面的优势',
                '__label__PATENTS':'优化文书中该项工在专利奖项方面的优势',
                '__label__INVESTMENT':'优化文书中该项工在投融资方面的优势',
                '__label__BUSINESS':'优化文书中该项目在商业市场的优势',
                '__label__INDUSTRIAL':'优化文书中该项在产业方面的优势',}
        if('FILES' in js.keys()):
            files = js['FILES']
            for f in files:
                fname = get_fname_by_path(f_path)
                if(fname!=f['NAME']):
                    continue
                labels = f['LABELS']
                for label in labels.keys():
                    if(labels[label]['score']==0):
                        txt=txt+ch_l[label]+'\n'

        for index,page in enumerate(pdf):

            if(index==0):


                page.addTextAnnot ((200,200), txt)
    fpath,fname=os.path.split(f_path)
    if not os.path.exists(fpath+'/res'):  #判断是否存在文件夹如果不存在则创建为文件夹
        os.mkdir(fpath+'/res')

    pdf.save(fpath+'/res/'+fname)


def get_res_files(folder_path,s,js):
    list = listdir(folder_path,[])
    for f in list :
        if(f.endswith('pdf')):
            pdf_hl(f,s,js)


